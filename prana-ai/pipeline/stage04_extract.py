"""
Stage 04 — LLM Extraction (manifest-driven)

Flow:
  1. Validate file format against manifest.supported_formats
  2. Multi-format OCR: PDF → PyMuPDF, DOCX → python-docx, images → Tesseract,
     XLSX → openpyxl, fallback → AWS Textract
  3. Redact NIK ([NIK_REDACTED]) before LLM sees text — pan_token computed in Stage 02
  4. If doc_type known: resolve manifest → build dynamic prompt → LLM extraction
  5. If doc_type is AUTO_DETECT:
       a. Run a lightweight probe extraction for indicator fields
       b. Score all manifests → pick best fit above AUTO_DETECT_MIN_SCORE
       c. Re-run full extraction with winning manifest
       d. If no manifest wins → return AutoDetectFailed (caller routes to unclassified_queue)
  6. Validate extraction confidence against manifest.confidence_threshold
  7. Return Stage04Result

Privacy: NIK is replaced with [NIK_REDACTED] in OCR text before LLM call.
         Raw PAN never reaches the LLM. pan_token was computed in Stage 02.
"""

import io
import json
import re
import logging
from dataclasses import dataclass, field
from typing import Optional

from extraction.extraction_service import ExtractionService, DocType
from llm_client import LLMClient
from manifest.manifest_client import ManifestClient, ManifestData
from manifest.prompt_builder import build_prompt, build_auto_detect_prompt
from pipeline.errors import PipelineError, PipelineException

log = logging.getLogger(__name__)

# Indian PAN pattern — redacted before LLM sees text
_PAN_RE = re.compile(r"\b[A-Z]{5}[0-9]{4}[A-Z]\b")

AUTO_DETECT_SENTINEL = "AUTO_DETECT"
AUTO_DETECT_MIN_SCORE = 0.5


@dataclass
class Stage04Result:
    extracted_fields:     dict
    overall_confidence:   float
    low_confidence_fields: list[str]
    doc_type:             str            # resolved doc_type (may differ from input if AUTO_DETECT)
    manifest_id:          str
    format_used:          str            # pdf | docx | image | xlsx | textract
    auto_detected:        bool = False   # True if doc_type was AUTO_DETECTed


@dataclass
class AutoDetectFailed:
    """Returned when AUTO_DETECT finds no matching manifest above threshold."""
    best_guess_doc_type:  Optional[str]
    best_guess_score:     float
    partial_fields:       dict           # indicator fields for OA-Admin context
    ocr_text_redacted:    str


class Stage04Extract:

    def __init__(
        self,
        llm_client: LLMClient,
        manifest_client: ManifestClient,
        aws_region: str = "ap-south-1",
    ):
        self._llm = llm_client
        self._manifests = manifest_client
        self._legacy_svc = ExtractionService(llm_client)   # fallback if manifest unavailable
        self._aws_region = aws_region

    async def run(
        self,
        file_bytes: bytes,
        ext: str,
        doc_type: str,           # declared by HRMS/OA, or AUTO_DETECT
        tenant_id: str,
        doc_period: Optional[str] = None,
    ) -> Stage04Result | AutoDetectFailed:

        ext = ext.lower().lstrip(".")

        # ── 1. OCR ──────────────────────────────────────────────────────────────
        raw_text, format_used = self._ocr(file_bytes, ext)  # raises PipelineException on total failure
        redacted_text = _PAN_RE.sub("[NIK_REDACTED]", raw_text)

        # Quality gate — blank output is non-retryable (document has no extractable text)
        if not redacted_text.strip() or len(redacted_text.strip()) < 20:
            raise PipelineException(
                code=PipelineError.S04_EXTRACT_OCR_BLANK_OUTPUT,
                stage="stage04",
                message=f"OCR produced blank or near-blank output ({len(redacted_text)} chars) for .{ext}",
            )

        # ── 2. AUTO_DETECT path ─────────────────────────────────────────────────
        if doc_type.upper() == AUTO_DETECT_SENTINEL:
            return await self._auto_detect(
                tenant_id, file_bytes, ext, redacted_text, format_used
            )

        # ── 3. Declared doc_type path ───────────────────────────────────────────
        try:
            manifest = await self._manifests.resolve(tenant_id, doc_type.upper())
        except ValueError:
            log.warning("No manifest for doc_type=%s, falling back to legacy extraction", doc_type)
            return await self._legacy_extract(doc_type, redacted_text, format_used)

        # Validate format
        if not manifest.format_supported(ext):
            log.warning(
                "Format %s not supported for doc_type=%s (supported: %s)",
                ext, doc_type, manifest.supported_formats,
            )
            return AutoDetectFailed(
                best_guess_doc_type=doc_type,
                best_guess_score=0.0,
                partial_fields={},
                ocr_text_redacted=redacted_text,
            )

        return await self._extract_with_manifest(manifest, redacted_text, format_used, auto_detected=False)

    # ── AUTO_DETECT ─────────────────────────────────────────────────────────────

    async def _auto_detect(
        self,
        tenant_id: str,
        file_bytes: bytes,
        ext: str,
        redacted_text: str,
        format_used: str,
    ) -> Stage04Result | AutoDetectFailed:

        # Step a: probe extraction for indicator fields
        probe_system, probe_user = build_auto_detect_prompt(redacted_text)
        probe_raw = await self._llm.complete(
            system=probe_system, user=probe_user, temperature=0.0
        )
        partial_fields = _safe_parse_json(probe_raw)
        # Flatten to {field: value} for scoring
        flat_partial = {
            k: v.get("value") if isinstance(v, dict) else v
            for k, v in partial_fields.items()
            if k != "overall_confidence"
        }

        # Step b: score all manifests
        all_manifests = await self._manifests.list_all(tenant_id)
        format_compatible = [m for m in all_manifests if m.format_supported(ext)]

        scored = sorted(
            [(m.score_against(flat_partial), m) for m in format_compatible],
            key=lambda x: x[0],
            reverse=True,
        )

        if not scored or scored[0][0] < AUTO_DETECT_MIN_SCORE:
            best_score, best_manifest = (scored[0] if scored else (0.0, None))
            log.info(
                "AUTO_DETECT failed: best=%s score=%.2f tenant=%s",
                best_manifest.doc_type if best_manifest else "none", best_score, tenant_id,
            )
            return AutoDetectFailed(
                best_guess_doc_type=best_manifest.doc_type if best_manifest else None,
                best_guess_score=best_score,
                partial_fields=flat_partial,
                ocr_text_redacted=redacted_text,
            )

        best_score, winning_manifest = scored[0]
        log.info(
            "AUTO_DETECT: classified as %s (score=%.2f) tenant=%s",
            winning_manifest.doc_type, best_score, tenant_id,
        )

        # Step c: full extraction with winning manifest
        return await self._extract_with_manifest(
            winning_manifest, redacted_text, format_used, auto_detected=True
        )

    # ── Manifest-driven extraction ──────────────────────────────────────────────

    async def _extract_with_manifest(
        self,
        manifest: ManifestData,
        redacted_text: str,
        format_used: str,
        auto_detected: bool,
    ) -> Stage04Result:
        system, user = build_prompt(manifest, redacted_text)
        try:
            raw = await self._llm.complete(system=system, user=user, temperature=0.0)
        except TimeoutError as exc:
            raise PipelineException(
                code=PipelineError.S04_EXTRACT_LLM_TIMEOUT,
                stage="stage04",
                message=f"LLM call timed out for {manifest.doc_type}: {exc}",
            ) from exc
        except Exception as exc:
            raise PipelineException(
                code=PipelineError.S04_EXTRACT_LLM_UNAVAILABLE,
                stage="stage04",
                message=f"LLM call failed for {manifest.doc_type}: {exc}",
            ) from exc
        parsed = _safe_parse_json(raw)
        if not parsed:
            raise PipelineException(
                code=PipelineError.S04_EXTRACT_LLM_JSON_INVALID,
                stage="stage04",
                message=f"LLM response unparseable as JSON for {manifest.doc_type}",
            )

        overall_confidence = float(parsed.get("overall_confidence", 0.0))

        # Compute low-confidence fields (required fields below manifest threshold)
        low_conf = [
            f for f in manifest.required_fields
            if isinstance(parsed.get(f), dict)
            and parsed[f].get("confidence", 0.0) < manifest.confidence_threshold
        ]

        log.info(
            "extraction complete: doc_type=%s confidence=%.2f low_conf=%s manifest=%s",
            manifest.doc_type, overall_confidence, low_conf, manifest.manifest_id,
        )

        return Stage04Result(
            extracted_fields=parsed,
            overall_confidence=overall_confidence,
            low_confidence_fields=low_conf,
            doc_type=manifest.doc_type,
            manifest_id=manifest.manifest_id,
            format_used=format_used,
            auto_detected=auto_detected,
        )

    async def _legacy_extract(
        self, doc_type: str, redacted_text: str, format_used: str
    ) -> Stage04Result:
        """Fallback to hardcoded prompt when manifest unavailable (e.g. dev/test)."""
        try:
            dt = DocType(doc_type.upper())
        except ValueError:
            dt = DocType.SALARY_SLIP  # safest default for unknown types
        result = await self._legacy_svc.extract(dt, redacted_text)
        return Stage04Result(
            extracted_fields=result.fields,
            overall_confidence=result.overall_confidence,
            low_confidence_fields=result.low_confidence_fields,
            doc_type=doc_type,
            manifest_id="legacy",
            format_used=format_used,
            auto_detected=False,
        )

    # ── Multi-format OCR ────────────────────────────────────────────────────────

    def _ocr(self, file_bytes: bytes, ext: str) -> tuple[str, str]:
        """
        Extract text from any supported format.
        Returns (text, format_used).

        Supported: pdf, docx, doc, jpeg, jpg, png, tiff, bmp, gif, xlsx, xls
        Fallback:  AWS Textract (handles most formats including scanned images)
        """
        try:
            if ext == "pdf":
                return self._ocr_pdf(file_bytes), "pdf"
            if ext in ("docx", "doc"):
                return self._ocr_docx(file_bytes), "docx"
            if ext in ("jpeg", "jpg", "png", "tiff", "tif", "bmp", "gif", "webp"):
                return self._ocr_image(file_bytes), "image"
            if ext in ("xlsx", "xls"):
                return self._ocr_xlsx(file_bytes), "xlsx"
            # Unknown extension — try Textract which handles many formats
            log.warning("Unknown extension %s — attempting Textract", ext)
            return self._textract_ocr(file_bytes), "textract"
        except PipelineException:
            raise  # password-protected and other structured errors propagate immediately
        except Exception as primary_err:
            log.warning("Primary OCR failed (%s), falling back to Textract: %s", ext, primary_err)
            try:
                return self._textract_ocr(file_bytes), "textract"
            except Exception as fallback_err:
                log.error("Textract fallback also failed: %s", fallback_err)
                raise PipelineException(
                    code=PipelineError.S04_EXTRACT_OCR_TEXTRACT_UNAVAILABLE,
                    stage="stage04",
                    message=f"Both primary OCR and Textract failed for .{ext}: {fallback_err}",
                ) from fallback_err

    def _ocr_pdf(self, file_bytes: bytes) -> str:
        """PDF → text via PyMuPDF (high quality, handles native + scanned PDFs)."""
        import fitz  # PyMuPDF
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as exc:
            msg = str(exc).lower()
            if "password" in msg or "encrypt" in msg:
                raise PipelineException(
                    code=PipelineError.S04_EXTRACT_PASSWORD_PROTECTED,
                    stage="stage04",
                    message="PDF is password-protected — cannot extract without employee-provided password",
                ) from exc
            raise
        if doc.needs_pass:
            raise PipelineException(
                code=PipelineError.S04_EXTRACT_PASSWORD_PROTECTED,
                stage="stage04",
                message="PDF requires a password to open",
            )
        pages_text = []
        for page in doc:
            # Try native text extraction first (fast, perfect for digitally generated PDFs)
            text = page.get_text("text")
            if len(text.strip()) < 50:
                # Low native text → scanned page → use Tesseract on rasterised image
                pix = page.get_pixmap(dpi=200)
                text = self._tesseract_pix(pix)
            pages_text.append(text)
        return "\n\n".join(pages_text)

    def _ocr_docx(self, file_bytes: bytes) -> str:
        """DOCX → text via python-docx. Preserves paragraph order."""
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        # Also extract table cells (common in salary slips exported as DOCX)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    def _ocr_image(self, file_bytes: bytes) -> str:
        """JPEG/PNG/TIFF/BMP → text via Tesseract (eng+hin)."""
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        # Upscale small images — Tesseract accuracy degrades below 300dpi equivalent
        min_dim = 1200
        if max(img.size) < min_dim:
            scale = min_dim / max(img.size)
            img = img.resize(
                (int(img.width * scale), int(img.height * scale)),
                Image.LANCZOS,
            )
        return pytesseract.image_to_string(img, lang="eng+hin")

    def _ocr_xlsx(self, file_bytes: bytes) -> str:
        """
        XLSX → text by reading all cell values.
        Common for bank statements and salary data from payroll software.
        """
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell).strip() for cell in row if cell is not None and str(cell).strip()
                )
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    def _tesseract_pix(self, pix) -> str:
        """Run Tesseract on a PyMuPDF Pixmap (used for scanned PDF pages)."""
        import pytesseract
        from PIL import Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        return pytesseract.image_to_string(img, lang="eng+hin")

    def _textract_ocr(self, file_bytes: bytes) -> str:
        """AWS Textract fallback — handles most formats including scanned images."""
        import boto3
        client = boto3.client("textract", region_name=self._aws_region)
        resp = client.detect_document_text(Document={"Bytes": file_bytes})
        lines = [
            b["Text"] for b in resp.get("Blocks", [])
            if b["BlockType"] == "LINE"
        ]
        return "\n".join(lines)


def _safe_parse_json(raw: str) -> dict:
    """Parse JSON from LLM response, stripping markdown fences if present."""
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        import re
        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw, re.DOTALL)
        if match:
            return json.loads(match.group(1))
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        log.error("LLM returned unparseable JSON: %s", raw[:300])
        return {}
